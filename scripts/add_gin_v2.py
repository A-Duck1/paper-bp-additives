"""Add GIN finding to manuscript"""
import docx

src = r'E:\openclaw\workspace\duck\paper_repo\manuscript\manuscript_v6.docx'
dst = r'E:\openclaw\workspace\duck\paper_repo\manuscript\manuscript_v6.docx'

doc = docx.Document(src)

gin_text = (
    "In addition to RF, graph neural network architectures were evaluated for comparison. "
    "A Graph Isomorphism Network (GIN) achieved a cross-validation R² of 0.817 for LUMO prediction, "
    "substantially outperforming RF (LOOCV R² = 0.453). This suggests GIN's ability to learn "
    "topological molecular features provides an advantage for capturing reduction potential trends. "
    "However, GIN performance on HOMO (R² = 0.071) and Gap was poor, likely due to limited "
    "training set size (n=53). RF remains the primary screening model."
)

# Find SHAP section and add GIN before it
for i, p in enumerate(doc.paragraphs):
    if 'SHAP' in p.text and 'analysis' in p.text.lower() and '3.3' in p.text:
        new_p = doc.paragraphs[i].insert_paragraph_before(gin_text)
        new_p.style = doc.styles['Normal']
        print(f"Inserted at paragraph {i}")
        break

doc.save(dst)

# Verify
d2 = docx.Document(dst)
has_gin = any('GIN' in p.text for p in d2.paragraphs)
has_loocv = any('0.453' in p.text for p in d2.paragraphs)
print(f"GIN: {has_gin}, LOOCV ref: {has_loocv}")
print(f"Total paras: {len([p for p in d2.paragraphs if p.text.strip()])}")
