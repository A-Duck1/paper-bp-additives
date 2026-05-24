"""Generate Cover Letter for Digital Discovery submission"""
import docx
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)

# Date
doc.add_paragraph('May 24, 2026')
doc.add_paragraph('')

# Editor info
doc.add_paragraph('Professor Alán Aspuru-Guzik')
doc.add_paragraph('Editor-in-Chief, Digital Discovery')
doc.add_paragraph('Royal Society of Chemistry')
doc.add_paragraph('')

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Dear Professor Aspuru-Guzik,')
run.bold = True

doc.add_paragraph('')
doc.add_paragraph(
    'We are pleased to submit our manuscript entitled '
    '"Machine Learning-Guided Screening of Boron/Phosphorus-Containing '
    'Bifunctional Additives for Synergistic SEI/CEI Regulation in '
    'Lithium-Ion Batteries" for consideration for publication in Digital Discovery.'
)
doc.add_paragraph('')

doc.add_paragraph(
    'The development of high-voltage lithium-ion batteries critically depends on '
    'the simultaneous stabilization of both the anode solid-electrolyte interphase (SEI) '
    'and the cathode-electrolyte interphase (CEI). While boron-containing and '
    'phosphorus-containing additives have been independently explored for SEI and CEI '
    'modification, respectively, no systematic screening framework has been reported for '
    'identifying bifunctional additives that can simultaneously stabilize both interfaces.'
)
doc.add_paragraph('')

doc.add_paragraph(
    'In this work, we present a machine learning framework for high-throughput screening of boron/phosphorus-containing bifunctional '
    'electrolyte additives. Key highlights include:'
)

# Bullet points
bullets = [
    'A curated dataset of 100 B/P-containing additives with experimentally known HOMO/LUMO energy levels, '
    'the most comprehensive open dataset for additive electronic properties to date.',
    'Random Forest models achieving R² = 0.85/0.83/0.80 for HOMO/LUMO/Gap prediction, '
    'with rigorous LOOCV validation (R² = 0.57/0.44/0.62) and SHAP-based chemical interpretability.',
    'A virtual screening of 14,425 candidate molecules across 152 chemical scaffolds, '
    'identifying top candidates with pyrophosphate and phosphonate moieties as promising '
    'dual SEI/CEI formers.',
    'Complete open-source code and data availability via GitHub (https://github.com/A-Duck1/paper-bp-additives), '
    'ensuring full reproducibility.',
]
for b in bullets:
    p = doc.add_paragraph(b, style='List Bullet')

doc.add_paragraph('')
doc.add_paragraph(
    'We believe this work aligns well with the scope of Digital Discovery, as it presents '
    'a reproducible computational workflow for accelerated materials discovery in energy storage. '
    'The combination of machine learning prediction, SHAP-based feature interpretation, and '
    'large-scale virtual screening provides a template for future data-driven electrolyte design.'
)
doc.add_paragraph('')
doc.add_paragraph(
    'This manuscript is original, has not been published previously, and is not under '
    'consideration elsewhere. All authors have approved the submission.'
)
doc.add_paragraph('')
doc.add_paragraph('We look forward to your response.')
doc.add_paragraph('')
doc.add_paragraph('Sincerely,')
doc.add_paragraph('Author Name')
doc.add_paragraph('Corresponding Author')
doc.add_paragraph('Email: name@institute.edu')

OUT = r".\paper_repo\manuscript\cover_letter.docx"
doc.save(OUT)
print(f"Cover letter: {OUT}")
