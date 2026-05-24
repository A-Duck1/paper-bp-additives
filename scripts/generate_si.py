#!/usr/bin/env python3
"""Generate Supplementary Information DOCX for the paper."""

import csv, json, os, sys
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ── Paths ──────────────────────────────────────────────────
BASE = os.path.dirname(os.path.abspath(__file__))
MANUSCRIPT_DIR = os.path.join(BASE, "..", "paper_repo", "manuscript")
DATA_DIR = os.path.join(BASE, "..", "paper_repo", "data")
ML_DIR = os.path.join(DATA_DIR, "ml_results")
OUTPUT = os.path.join(MANUSCRIPT_DIR, "supplementary_information.docx")

TRAINING_V5 = r"D:\pubchem_data\training_v5.csv"
TRAINING_V4 = r"D:\pubchem_data\training_v4.csv"
TOP200 = os.path.join(ML_DIR, "screening_rf_v7_top200.csv")
SHAP = os.path.join(ML_DIR, "shap_analysis.json")
XGB_LOOCV = os.path.join(ML_DIR, "xgb_loocv_results.json")

os.makedirs(MANUSCRIPT_DIR, exist_ok=True)


# ── Helpers ─────────────────────────────────────────────────
def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color,
    })
    shading.append(shading_elem)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a table with header styling."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = str(h)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
        set_cell_shading(cell, "D9E2F3")

    # Data rows
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val) if val is not None else ""
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(8.5)

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)

    return table


def load_training_data():
    """Load training data from v5 if exists, else v4."""
    path = TRAINING_V5 if os.path.exists(TRAINING_V5) else TRAINING_V4
    rows = []
    with open(path, "r", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for r in reader:
            rows.append(r)
    return rows, path


def load_top200():
    """Load top 200 screening results."""
    rows = []
    with open(TOP200, "r", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for r in reader:
            rows.append(r)
    return rows


def load_shap():
    """Load SHAP analysis."""
    with open(SHAP, "r", encoding="utf-8") as f:
        return json.load(f)


def load_xgb_loocv():
    """Load XGB LOOCV results."""
    with open(XGB_LOOCV, "r", encoding="utf-8") as f:
        return json.load(f)


# ── Document Builder ───────────────────────────────────────
def build_document():
    doc = Document()

    # ── Title ──
    title = doc.add_heading("Supplementary Information", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "This document provides supplementary data and analysis results "
        "supporting the main manuscript. All tables and figures are referenced "
        "in the main text as Section S1–S5."
    )

    # =========================================================
    # SECTION S1: Full Dataset Composition
    # =========================================================
    doc.add_heading("Section S1: Full Dataset Composition", level=1)
    
    training_data, training_path = load_training_data()
    n_train = len(training_data)
    dataset_name = os.path.basename(training_path)
    
    doc.add_paragraph(
        f"Table S1 lists all {n_train} training additives used in this study "
        f"(source: {dataset_name}). "
        "Additives are categorized by type: Boron-based (B), "
        "Phosphorus-based (P), Boron-Phosphorus (BP), and Reference (Ref)."
    )

    # Build the training table
    # Determine columns from first row
    sample_row = training_data[0]
    
    # For v4: name, smiles, homo, lumo, gap, type
    # For v5: name, smiles, homo, lumo, gap, category, source
    category_col = "category" if "category" in sample_row else "type"
    
    headers = ["#", "Name", "SMILES", "Category", "HOMO (eV)", "LUMO (eV)", "Gap (eV)"]
    rows_data = []
    for i, r in enumerate(training_data, 1):
        rows_data.append([
            i,
            r.get("name", ""),
            r.get("smiles", ""),
            r.get(category_col, ""),
            round(float(r["homo"]), 3) if r.get("homo") else "",
            round(float(r["lumo"]), 3) if r.get("lumo") else "",
            round(float(r["gap"]), 3) if r.get("gap") else "",
        ])
    
    caption = doc.add_paragraph(f"Table S1. Complete list of {n_train} training additives.")
    caption.runs[0].bold = True
    
    add_styled_table(doc, headers, rows_data,
                     col_widths=[1.0, 3.5, 5.0, 2.0, 2.5, 2.5, 2.5])

    doc.add_paragraph().add_run("").font.size = Pt(6)  # spacing

    # Category summary
    doc.add_paragraph("Category distribution:")
    from collections import Counter
    cats = Counter(r.get(category_col, "") for r in training_data)
    for cat, cnt in sorted(cats.items()):
        doc.add_paragraph(f"• {cat}: {cnt} samples", style="List Bullet")

    # =========================================================
    # SECTION S2: Virtual Screening Full Results (Top 50)
    # =========================================================
    doc.add_heading("Section S2: Virtual Screening Full Results", level=1)
    doc.add_paragraph(
        "Table S2 presents the top 50 candidate molecules identified through "
        "RF-based virtual screening of 14,425 PubChem candidates. "
        "Ranking is based on the composite score combining normalized HOMO, "
        "LUMO, and Gap predictions."
    )

    top200 = load_top200()
    top50 = top200[:50]

    headers = [
        "Rank", "SMILES", "RF_HOMO (eV)", "RF_LUMO (eV)", "RF_Gap (eV)",
        "Score", "MW (g/mol)"
    ]
    rows_data = []
    for r in top50:
        rows_data.append([
            r.get("rank", ""),
            r.get("smiles", ""),
            round(float(r["RF_homo"]), 3) if r.get("RF_homo") else "",
            round(float(r["RF_lumo"]), 3) if r.get("RF_lumo") else "",
            round(float(r["RF_gap"]), 3) if r.get("RF_gap") else "",
            round(float(r["score"]), 4) if r.get("score") else "",
            round(float(r["MW"]), 1) if r.get("MW") else "",
        ])

    caption = doc.add_paragraph("Table S2. Top 50 candidate molecules from RF-based virtual screening.")
    caption.runs[0].bold = True

    add_styled_table(doc, headers, rows_data,
                     col_widths=[1.2, 5.5, 2.5, 2.5, 2.5, 2.0, 2.5])

    doc.add_paragraph().add_run("").font.size = Pt(6)

    top200_stats_path = os.path.join(ML_DIR, "screening_rf_v7_stats.json")
    if os.path.exists(top200_stats_path):
        with open(top200_stats_path, "r") as f:
            stats = json.load(f)
        doc.add_paragraph(f"Summary statistics (Top 200):")
        doc.add_paragraph(f"• Score range: {stats['score']['min']:.4f} – {stats['score']['max']:.4f} "
                          f"(mean: {stats['score']['mean']:.4f} ± {stats['score']['std']:.4f})")
        doc.add_paragraph(f"• HOMO range: {stats['homo']['min']:.3f} – {stats['homo']['max']:.3f} eV")
        doc.add_paragraph(f"• LUMO range: {stats['lumo']['min']:.3f} – {stats['lumo']['max']:.3f} eV")
        doc.add_paragraph(f"• Gap range: {stats['gap']['min']:.3f} – {stats['gap']['max']:.3f} eV")
        doc.add_paragraph(f"• Candidate types in top 200: "
                          f"{stats['top200']['n_B']} B, {stats['top200']['n_P']} P, "
                          f"{stats['top200']['n_BP']} BP "
                          f"(MW: {stats['top200']['MW_mean']:.1f} ± range "
                          f"{stats['top200']['MW_min']:.1f}–{stats['top200']['MW_max']:.1f})")

    # =========================================================
    # SECTION S3: SHAP Feature Importance (Full)
    # =========================================================
    doc.add_heading("Section S3: SHAP Feature Importance (Full Top 30)", level=1)
    doc.add_paragraph(
        "Tables S3–S5 list the top 30 most important features identified by "
        "SHAP analysis for HOMO, LUMO, and Gap predictions, respectively. "
        "Importance is measured as mean |SHAP| value across all training samples."
    )

    shap_data = load_shap()
    results = shap_data["results"]
    target_names = {"homo": "S3 (HOMO)", "lumo": "S4 (LUMO)", "gap": "S5 (Gap)"}
    target_labels = {"homo": "HOMO", "lumo": "LUMO", "gap": "Gap"}

    for target in ["homo", "lumo", "gap"]:
        features = results[target]
        label = target_labels[target]
        table_id = target_names[target]

        caption = doc.add_paragraph(
            f"Table {table_id}. Top {len(features)} features for {label} prediction."
        )
        caption.runs[0].bold = True

        headers = ["Rank", "Feature", "Mean |SHAP|", "Mean SHAP"]
        rows_data = []
        for f in features:
            rows_data.append([
                f["rank"],
                f["feature"],
                f"{f['mean_abs_shap']:.5f}",
                f"{f['mean_shap']:.5f}",
            ])

        add_styled_table(doc, headers, rows_data,
                         col_widths=[1.5, 5.0, 3.0, 3.0])
        doc.add_paragraph().add_run("").font.size = Pt(6)

    # =========================================================
    # SECTION S4: LOOCV Detailed Results
    # =========================================================
    doc.add_heading("Section S4: Leave-One-Out Cross-Validation Results", level=1)
    doc.add_paragraph(
        "Table S6 presents the complete LOOCV results for both Random Forest "
        "(RF) and XGBoost (XGB) models across three target properties (HOMO, "
        "LUMO, Gap) and two feature sets (217 RDKit descriptors and 2048-bit "
        "Morgan fingerprints)."
    )

    # RF LOOCV (descriptors) from metrics_v4.json
    METRICS_V4 = os.path.join(BASE, "..", "data", "ml_results", "metrics_v4.json")
    rf_desc = {}
    if os.path.exists(METRICS_V4):
        with open(METRICS_V4) as f:
            rf_desc = json.load(f)

    # RF LOOCV (Morgan) - from reference in xgb_loocv.py
    rf_morgan = {
        "homo": {"r2": 0.567, "mae": None},
        "lumo": {"r2": 0.435, "mae": None},
        "gap": {"r2": 0.622, "mae": None},
    }

    # XGB LOOCV
    xgb_data = load_xgb_loocv()
    xgb_desc = xgb_data.get("descriptors", {})
    xgb_morgan = xgb_data.get("morgan", {})

    # Build the combined LOOCV table
    headers = [
        "Target",
        "Model", "Feature Set",
        "LOOCV R²", "LOOCV MAE"
    ]
    rows_data = []

    targets = ["homo", "lumo", "gap"]
    target_labels_short = {"homo": "HOMO", "lumo": "LUMO", "gap": "Gap"}

    for t in targets:
        label = target_labels_short[t]

        # RF descriptors
        if t in rf_desc:
            rows_data.append([
                label, "RF", "217 RDKit",
                f"{rf_desc[t]['r2']:.4f}",
                f"{rf_desc[t]['mae']:.4f}" if rf_desc[t].get("mae") is not None else "N/A",
            ])

        # RF Morgan
        if t in rf_morgan:
            rows_data.append([
                label, "RF", "2048-bit Morgan",
                f"{rf_morgan[t]['r2']:.4f}",
                f"{rf_morgan[t]['mae']:.4f}" if rf_morgan[t].get("mae") is not None else "N/A",
            ])

        # XGB descriptors
        if t in xgb_desc:
            rows_data.append([
                label, "XGB", "217 RDKit",
                f"{xgb_desc[t]['loocv_r2']:.4f}",
                f"{xgb_desc[t]['loocv_mae']:.4f}",
            ])

        # XGB Morgan
        if t in xgb_morgan:
            rows_data.append([
                label, "XGB", "2048-bit Morgan",
                f"{xgb_morgan[t]['loocv_r2']:.4f}",
                f"{xgb_morgan[t]['loocv_mae']:.4f}",
            ])

    caption = doc.add_paragraph(
        "Table S6. Complete LOOCV results: RF vs XGB across all targets and feature sets."
    )
    caption.runs[0].bold = True

    add_styled_table(doc, headers, rows_data,
                     col_widths=[2.0, 2.0, 4.0, 3.0, 3.0])

    doc.add_paragraph().add_run("").font.size = Pt(6)

    # Summary paragraph
    avg_r2_rf_desc = sum(rf_desc[t]["r2"] for t in targets if t in rf_desc) / 3
    avg_r2_rf_morgan = sum(rf_morgan[t]["r2"] for t in targets if t in rf_morgan) / 3
    avg_r2_xgb_desc = sum(
        xgb_desc[t]["loocv_r2"]
        for t in targets if t in xgb_desc
    ) / 3
    avg_r2_xgb_morgan = sum(
        xgb_morgan[t]["loocv_r2"]
        for t in targets if t in xgb_morgan
    ) / 3

    doc.add_paragraph(
        f"Overall, RF with RDKit descriptors achieved the highest average "
        f"LOOCV R² ({avg_r2_rf_desc:.4f}), followed by XGB with RDKit descriptors "
        f"({avg_r2_xgb_desc:.4f}), RF with Morgan fingerprints ({avg_r2_rf_morgan:.4f}), "
        f"and XGB with Morgan fingerprints ({avg_r2_xgb_morgan:.4f})."
    )

    # =========================================================
    # SECTION S5: Model Hyperparameters
    # =========================================================
    doc.add_heading("Section S5: Model Hyperparameters", level=1)
    doc.add_paragraph(
        "Tables S7 and S8 summarize the hyperparameters used for the RF and "
        "XGBoost models in this study."
    )

    # RF
    caption = doc.add_paragraph("Table S7. Random Forest hyperparameters.")
    caption.runs[0].bold = True
    rf_headers = ["Parameter", "Value"]
    rf_params = [
        ("n_estimators", "500"),
        ("max_depth", "15"),
        ("min_samples_split", "5"),
        ("min_samples_leaf", "2"),
        ("random_state", "42"),
        ("n_jobs", "-1 (all cores)"),
    ]
    add_styled_table(doc, rf_headers, rf_params, col_widths=[5.0, 5.0])

    doc.add_paragraph().add_run("").font.size = Pt(6)

    # XGB
    caption = doc.add_paragraph("Table S8. XGBoost hyperparameters.")
    caption.runs[0].bold = True
    xgb_headers = ["Parameter", "Value"]
    xgb_params = [
        ("n_estimators", "300"),
        ("max_depth", "8"),
        ("learning_rate", "0.08"),
        ("subsample", "0.8"),
        ("colsample_bytree", "0.8"),
        ("random_state", "42"),
    ]
    add_styled_table(doc, xgb_headers, xgb_params, col_widths=[5.0, 5.0])

    # ── Save ──
    doc.save(OUTPUT)
    print(f"✅ Supplementary Information saved to: {OUTPUT}")


if __name__ == "__main__":
    build_document()
