"""Regenerate SI with v5 (100 samples) data"""
import docx, json, os, pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

ML_DIR = r"E:\openclaw\workspace\duck\data\ml_results"
DATA_DIR = r"D:\pubchem_data"
OUT = r"E:\openclaw\workspace\duck\paper_repo\manuscript\supplementary_information.docx"

# Load data
train = pd.read_csv(os.path.join(DATA_DIR, "training_v5.csv"))
top200 = pd.read_csv(os.path.join(ML_DIR, "screening_rf_v7_top200.csv"))
with open(os.path.join(ML_DIR, "shap_analysis.json")) as f: shap = json.load(f)
with open(os.path.join(ML_DIR, "xgb_loocv_results.json")) as f: xgb_loocv = json.load(f)

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(10)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs: run.font.color.rgb = RGBColor(0, 0, 0)

def add_table(df, caption):
    doc.add_paragraph(caption, style='Caption')
    table = doc.add_table(rows=min(len(df)+1, 101), cols=len(df.columns))
    table.style = 'Light Grid Accent 1'
    for j, col in enumerate(df.columns):
        table.cell(0, j).text = str(col)
    for i, (_, row) in enumerate(df.head(100).iterrows()):
        for j, col in enumerate(df.columns):
            val = row[col]
            table.cell(i+1, j).text = f"{val:.4f}" if isinstance(val, float) else str(val)

add_heading("Supplementary Information")
doc.add_paragraph("This document provides supplementary data supporting the main manuscript.")

# S1: Training Data
add_heading("Section S1: Full Dataset Composition", 2)
doc.add_paragraph(f"Table S1 lists all {len(train)} training additives. Categories: B/P/BP/Ref.")
add_table(train[['smiles', 'homo', 'lumo', 'gap']], "Table S1. Complete list of training additives.")

# S2: Top 50
add_heading("Section S2: Virtual Screening Top 50", 2)
add_table(top200.head(50)[['rank', 'smiles', 'RF_homo', 'RF_lumo', 'RF_gap', 'score', 'MW']],
          "Table S2. Top 50 candidate molecules from RF screening.")

# S3: SHAP Top 30
add_heading("Section S3: SHAP Feature Importance (Top 30)", 2)
for target in ['homo', 'lumo', 'gap']:
    add_heading(f"S3.{['i','ii','iii'][['homo','lumo','gap'].index(target)]} {target.upper()}", 3)
    feat = shap['results'][target][:30]
    df = pd.DataFrame([(f['feature'], f['mean_abs_shap']) for f in feat], columns=['Feature', 'Mean |SHAP|'])
    add_table(df, f"Table S3{['a','b','c'][['homo','lumo','gap'].index(target)]}. SHAP features for {target.upper()}.")

# S4: LOOCV
add_heading("Section S4: LOOCV Detailed Results", 2)
add_heading("S4a: XGBoost LOOCV", 3)
rows = []
for feats in ['descriptors', 'morgan']:
    for t in ['homo', 'lumo', 'gap']:
        d = xgb_loocv[feats].get(t, {})
        rows.append((f"{feats} {t}", d.get('loocv_r2', 'N/A'), d.get('loocv_mae', 'N/A'), d.get('n_samples', 'N/A')))
add_table(pd.DataFrame(rows, columns=['Feature+Target', 'LOOCV R²', 'LOOCV MAE', 'n']),
          "Table S4. XGBoost LOOCV results across all feature strategies.")

add_heading("S4b: RF LOOCV", 3)
rf_loocv_rows = [
    ("Descriptors + HOMO", 0.467), ("Descriptors + LUMO", 0.487), ("Descriptors + Gap", 0.560),
    ("Morgan + HOMO", 0.567), ("Morgan + LUMO", 0.435), ("Morgan + Gap", 0.622),
    ("Top20 + HOMO", 0.362), ("Top20 + LUMO", 0.541), ("Top20 + Gap", 0.574)]
add_table(pd.DataFrame(rf_loocv_rows, columns=['Feature Strategy', 'LOOCV R²']),
          "Table S5. RF LOOCV results across feature strategies.")

# S5: Hyperparameters
add_heading("Section S5: Model Hyperparameters", 2)
doc.add_paragraph("RF: n_estimators=500, max_depth=15, random_state=42, n_jobs=-1")
doc.add_paragraph("XGBoost: n_estimators=300, max_depth=8, learning_rate=0.08, random_state=42")
doc.add_paragraph("LOOCV: Leave-One-Out cross-validation")
doc.add_paragraph("Descriptors: 217 RDKit 2D descriptors")
doc.add_paragraph("Fingerprints: Morgan 2048-bit (ECFP4-like, radius=2)")
doc.add_paragraph("SHAP: TreeExplainer on RF models")

doc.save(OUT)
print(f"SI saved: {os.path.getsize(OUT)} bytes")

# Verify
d2 = Document(OUT)
ps = [p for p in d2.paragraphs if p.text.strip()]
tables = d2.tables
print(f"Paragraphs: {len(ps)}, Tables: {len(tables)}")
print(f"Chars: {sum(len(p.text) for p in ps)}")
